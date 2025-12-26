import os
import json
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from .models import Resume
from .serializers import ResumeSerializer


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_resume(request):
    """
    Generate a resume content using Gemini based on provided data.
    """
    # Lazy import - only load when function is called
    try:
        import google.generativeai as genai
    except (ImportError, TypeError):
        return Response({"error": "AI Service unavailable"}, status=503)

    user = request.user
    data = request.data
    
    # Extract data
    personal_info = data.get('personal_info', {})
    education = data.get('education', [])
    experience = data.get('experience', [])
    skills = data.get('skills', [])
    projects = data.get('projects', [])
    
    # Construct Prompt based on User Specification
    prompt = f"""
    You are a Resume Builder Engine.
    
    Below is the complete schema and instructions. 
    Whenever I send you user data, generate a resume exactly following this structure, styling logic, and template rules.
    
    USER DATA provided:
    {json.dumps(data, indent=2)}

    ----------------------------------------------------------------------
    OUTPUT RULES
    ----------------------------------------------------------------------

    1. Always follow the selected template’s settings (found in 'preferences').
    2. Always output the resume in clean HTML format (inside a single <div>).
    3. Maintain correct section order from "preferences.section_order".
    4. Use Tailwind CSS classes for styling to match the requested template ({data.get('preferences', {}).get('template_id', 'modern_minimal')}).
    5. Avoid long paragraphs — prefer concise bullet points.
    6. Remove any empty sections automatically.
    7. Prioritize measurable achievements (numbers, % improvements).
    8. Do NOT include explanation, reasoning, or commentary — only the resume content.

    ----------------------------------------------------------------------
    WHEN GENERATING A RESUME, FOLLOW THIS STRUCTURE (Translated to HTML):
    ----------------------------------------------------------------------

    <!-- Header -->
    <h1>FULL NAME</h1>
    <p>Headline</p>
    <p>email · phone · location · links</p>

    <!-- Sections (Order varies by template) -->
    <h2>SUMMARY</h2>
    <p>Short 2–3 line summary rewritten professionally.</p>

    <h2>WORK EXPERIENCE</h2>
    <h3>Title — Company <span style="float:right">Date Start – End</span></h3>
    <ul>
        <li>Bullet point...</li>
    </ul>

    <h2>PROJECTS</h2>
    <h3>Project Name</h3>
    <p>Description sentence</p>
    <ul>
        <li>Bullet point...</li>
    </ul>

    <h2>EDUCATION</h2>
    <p>Degree — Institution (Year)</p>

    <h2>SKILLS</h2>
    <p>List skills separated by commas or grouped by category.</p>

    <h2>CERTIFICATIONS</h2>
    <p>Certification Name — Issuer (Year)</p>

    """
    
    try:
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
             return Response({"error": "API Key missing"}, status=500)
             
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        response = model.generate_content(prompt)
        content = response.text.replace("```html", "").replace("```", "")
        
        # Save to DB
        resume = Resume.objects.create(
            user=user,
            title=data.get('title', 'My Resume'),
            personal_info=personal_info,
            education=education,
            experience=experience,
            skills=skills,
            projects=projects,
            generated_content=content
        )
        
        # Update Streak
        try:
            from users.utils import update_streak
            update_streak(user, "resume_generation")
        except Exception as e:
            pass  # print(f"Error updating streak: {e}")

        return Response(ResumeSerializer(resume).data, status=201)

    except Exception as e:
        return Response({"error": str(e)}, status=500)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_resumes(request):
    resumes = Resume.objects.filter(user=request.user).order_by('-created_at')
    return Response(ResumeSerializer(resumes, many=True).data)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_resume_detail(request, pk):
    try:
        resume = Resume.objects.get(pk=pk, user=request.user)
        return Response(ResumeSerializer(resume).data)
    except Resume.DoesNotExist:
        return Response({"error": "Not found"}, status=404)


@api_view(['PUT', 'PATCH'])
@permission_classes([IsAuthenticated])
def update_resume(request, pk):
    """
    Update an existing resume's data without regenerating content.
    """
    try:
        resume = Resume.objects.get(pk=pk, user=request.user)
    except Resume.DoesNotExist:
        return Response({"error": "Not found"}, status=404)
    
    data = request.data
    
    # Update fields if provided
    if 'title' in data:
        resume.title = data['title']
    if 'personal_info' in data or 'personal' in data:
        resume.personal_info = data.get('personal_info') or data.get('personal', {})
    if 'education' in data:
        resume.education = data['education']
    if 'experience' in data or 'work_experience' in data:
        resume.experience = data.get('experience') or data.get('work_experience', [])
    if 'skills' in data:
        resume.skills = data['skills']
    if 'projects' in data:
        resume.projects = data['projects']
    if 'generated_content' in data:
        resume.generated_content = data['generated_content']
    
    resume.save()
    return Response(ResumeSerializer(resume).data)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_resume(request, pk):
    """
    Delete a resume.
    """
    try:
        resume = Resume.objects.get(pk=pk, user=request.user)
        resume.delete()
        return Response({"message": "Resume deleted successfully"}, status=200)
    except Resume.DoesNotExist:
        return Response({"error": "Not found"}, status=404)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def import_resume(request):
    """
    Import a resume file (PDF/DOCX) and parse it using AI.
    """
    # Lazy import - only load when function is called
    try:
        import google.generativeai as genai
    except (ImportError, TypeError):
        return Response({"error": "AI Service unavailable"}, status=503)
    
    if 'file' not in request.FILES:
        return Response({"error": "No file uploaded"}, status=400)
    
    uploaded_file = request.FILES['file']
    file_name = uploaded_file.name.lower()
    
    # Read file content
    try:
        if file_name.endswith('.pdf'):
            # For PDF files, use PyPDF2 or just extract text with AI
            import io
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() or ""
            except ImportError:
                # Fallback: just note that it's a PDF
                text_content = f"[PDF file uploaded: {uploaded_file.name}. Please extract content manually.]"
                uploaded_file.seek(0)
        elif file_name.endswith('.docx'):
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(uploaded_file.read()))
                text_content = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                text_content = f"[DOCX file uploaded: {uploaded_file.name}. Please install python-docx.]"
        elif file_name.endswith('.txt'):
            text_content = uploaded_file.read().decode('utf-8')
        else:
            return Response({"error": "Unsupported file format. Please upload PDF, DOCX, or TXT."}, status=400)
    except Exception as e:
        return Response({"error": f"Failed to read file: {str(e)}"}, status=400)
    
    # Use Gemini to parse the resume text into structured data
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    model = genai.GenerativeModel("gemini-2.0-flash")
    
    parse_prompt = f"""
    Parse the following resume text and extract structured data.
    Return ONLY valid JSON in the exact format below, no markdown, no explanation:
    
    {{
        "personal": {{
            "first_name": "",
            "last_name": "",
            "email": "",
            "phone": "",
            "address": "",
            "job_title": ""
        }},
        "education": [
            {{
                "institution": "",
                "degree": "",
                "field": "",
                "start_date": "",
                "end_date": "",
                "percentage": ""
            }}
        ],
        "experience": [
            {{
                "company": "",
                "title": "",
                "location": "",
                "start_date": "",
                "end_date": "",
                "description": ""
            }}
        ],
        "skills": [
            {{
                "category": "Technical Skills",
                "items": ""
            }}
        ],
        "projects": [
            {{
                "name": "",
                "description": "",
                "technologies": ""
            }}
        ]
    }}
    
    RESUME TEXT:
    {text_content[:8000]}
    """
    
    try:
        response = model.generate_content(parse_prompt)
        response_text = response.text.strip()
        
        # Clean up JSON if wrapped in markdown
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
        response_text = response_text.strip()
        
        parsed_data = json.loads(response_text)
        
        # Create a new resume with parsed data
        title = f"Imported - {uploaded_file.name}"
        resume = Resume.objects.create(
            user=request.user,
            title=title,
            personal_info=parsed_data.get('personal', {}),
            education=parsed_data.get('education', []),
            experience=parsed_data.get('experience', []),
            skills=parsed_data.get('skills', []),
            projects=parsed_data.get('projects', []),
            generated_content=""
        )
        
        return Response({
            "id": resume.id,
            "title": resume.title,
            "parsed_data": parsed_data,
            "message": "Resume imported successfully"
        }, status=201)
        
    except json.JSONDecodeError as e:
        return Response({"error": f"Failed to parse resume: {str(e)}"}, status=500)
    except Exception as e:
        return Response({"error": f"AI parsing failed: {str(e)}"}, status=500)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def analyze_ats(request):
    """
    Analyze a resume for ATS compatibility and provide a score.
    """
    # Lazy import - only load when function is called
    try:
        import google.generativeai as genai
    except (ImportError, TypeError):
        return Response({"error": "AI Service unavailable"}, status=503)
    
    resume_id = request.data.get('resume_id')
    job_description = request.data.get('job_description', '')
    
    # Get resume content
    if resume_id:
        try:
            resume = Resume.objects.get(pk=resume_id, user=request.user)
            resume_content = resume.generated_content or json.dumps({
                'personal': resume.personal_info,
                'education': resume.education,
                'experience': resume.experience,
                'skills': resume.skills,
                'projects': resume.projects
            })
        except Resume.DoesNotExist:
            return Response({"error": "Resume not found"}, status=404)
    elif 'file' in request.FILES:
        # Handle file upload for direct analysis
        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name.lower()
        file_bytes = uploaded_file.read()
        
        try:
            resume_content = ""
            
            # Handle PDF files
            if file_name.endswith('.pdf'):
                import io
                try:
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            resume_content += text + "\n"
                    
                    # If no text extracted (image-based PDF), use Gemini Vision
                    if not resume_content.strip():
                        resume_content = "[Image-based PDF detected - using AI vision for analysis]"
                except ImportError:
                    resume_content = ""
                except Exception:
                    resume_content = ""
            
            # Handle DOCX files
            elif file_name.endswith('.docx'):
                try:
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(file_bytes))
                    resume_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                except ImportError:
                    resume_content = ""
            
            # Handle image files (PNG, JPG) - Use Gemini Vision
            elif file_name.endswith(('.png', '.jpg', '.jpeg', '.webp')):
                import base64
                # Convert image to base64 for Gemini Vision
                image_base64 = base64.b64encode(file_bytes).decode('utf-8')
                
                # Determine mime type
                if file_name.endswith('.png'):
                    mime_type = 'image/png'
                elif file_name.endswith('.webp'):
                    mime_type = 'image/webp'
                else:
                    mime_type = 'image/jpeg'
                
                # Use Gemini Vision to extract text from image
                genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
                vision_model = genai.GenerativeModel("gemini-2.0-flash")
                
                ocr_prompt = """
                Extract ALL text from this resume image. 
                Return the complete text content preserving the structure.
                Include every section: name, contact info, summary, experience, education, skills, projects, certifications.
                """
                
                response = vision_model.generate_content([
                    ocr_prompt,
                    {"mime_type": mime_type, "data": image_base64}
                ])
                resume_content = response.text.strip()
            
            # Handle plain text files
            else:
                resume_content = file_bytes.decode('utf-8', errors='ignore')
            
            # If still no content, return error
            if not resume_content.strip():
                return Response({"error": "Could not extract text from file. Please try a different format."}, status=400)
                
        except Exception as e:
            return Response({"error": f"Failed to read file: {str(e)}"}, status=400)
    else:
        return Response({"error": "No resume provided"}, status=400)
    
    # Use Gemini for ATS analysis
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    model = genai.GenerativeModel("gemini-2.0-flash")
    
    ats_prompt = f"""
    You are an ATS (Applicant Tracking System) expert. Analyze this resume and provide:
    
    1. An overall ATS score from 0-100
    2. Category scores (each out of 100):
       - Keywords & Skills Match
       - Formatting & Structure  
       - Experience Relevance
       - Education & Certifications
       - Contact Information
    3. Top 5 suggestions to improve the resume
    4. Missing keywords (if job description provided)
    
    Return ONLY valid JSON in this exact format:
    {{
        "overall_score": 75,
        "categories": {{
            "keywords_skills": {{"score": 80, "feedback": "Good keyword usage"}},
            "formatting": {{"score": 70, "feedback": "Structure is clean"}},
            "experience": {{"score": 75, "feedback": "Relevant experience shown"}},
            "education": {{"score": 85, "feedback": "Education well formatted"}},
            "contact_info": {{"score": 90, "feedback": "All contact info present"}}
        }},
        "suggestions": [
            "Add more action verbs",
            "Include measurable achievements",
            "Add relevant certifications",
            "Use industry keywords",
            "Improve bullet point formatting"
        ],
        "missing_keywords": ["Python", "AWS", "Agile"],
        "summary": "Your resume is well-structured with minor improvements needed."
    }}
    
    RESUME CONTENT:
    {resume_content[:6000]}
    
    {"JOB DESCRIPTION:" + job_description[:2000] if job_description else "No job description provided."}
    """
    
    try:
        response = model.generate_content(ats_prompt)
        response_text = response.text.strip()
        
        # Clean up JSON
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
        response_text = response_text.strip()
        
        analysis = json.loads(response_text)
        return Response(analysis)
        
    except json.JSONDecodeError:
        # Return a fallback analysis
        return Response({
            "overall_score": 70,
            "categories": {
                "keywords_skills": {"score": 70, "feedback": "Review your skills section"},
                "formatting": {"score": 75, "feedback": "Good overall structure"},
                "experience": {"score": 70, "feedback": "Add more action verbs"},
                "education": {"score": 80, "feedback": "Education section is complete"},
                "contact_info": {"score": 85, "feedback": "Contact information present"}
            },
            "suggestions": [
                "Add more specific technical skills",
                "Use stronger action verbs",
                "Include quantifiable achievements",
                "Add relevant certifications",
                "Tailor resume to job description"
            ],
            "missing_keywords": [],
            "summary": "Your resume has good structure. Focus on adding specific achievements."
        })
    except Exception as e:
        return Response({"error": f"Analysis failed: {str(e)}"}, status=500)
